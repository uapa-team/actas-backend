from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField
from ..models import Request
from .case_utils import add_analysis_paragraph


class TEPR(Request):

    full_name = 'Tránsito entre programas'

    origin_program = StringField(
        min_length=4, max_length=4, choices=Request.PLAN_CHOICES,
        required=True, display='Programa Académico origen', default=Request.PI_ELECTRICA)
    academic_period_transit = StringField(
        max_length=10, required=True, display='Periodo de tránsito', 
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    available_places = BooleanField(
        display='Hay cupos en el plan de estudios', default=False)
    languaje = BooleanField(
        display='Cumple requisito de idioma en el plan destino', default=True)
    on_time = BooleanField(
        display='Cumple fechas establecidas', default=True)

    regulation_list = ['035|2014|VAC', '002|2011|CFA',
                       '241|2009|VAC']  # List of regulations

    str_cm = [
        'tránsito del programa {} ({}) al programa {} ({}), a partir del periodo académico {}',
        'debido a que {}.'
    ]

    str_pcm = [
        'En el programa hay {}cupos para tránsito.',
        'Viene del programa {} ({}).',
        'El estudiante {}cumple con la suficiencia de idioma exigida.',
        'La solicitud {}se hace luego de completar el plan de estudios y antes del grado (a menos' +
        ' que no se haya abierto convocatorio durante el periodo)(Parágrafo 2, {}).'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0].format(
                # pylint: disable=no-member
                self.get_origin_program_display(),
                self.origin_program,
                self.get_academic_program_display(),
                self.academic_program,
                self.academic_period_transit
            ) + ', '
        )
        paragraph.add_run(
            self.str_cm[1].format(
                self.council_decision
            )
        )

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            '' if self.available_places else 'no '
        )]
        analysis_list += [self.str_pcm[1].format(
            # pylint: disable=no-member
            self.get_origin_program_display(),
            self.origin_program,
        )]
        analysis_list += [self.str_pcm[2].format(
            '' if self.languaje else 'no '
        )]
        analysis_list += [self.str_pcm[3].format(
            '' if self.on_time else 'no ',
            Request.regulations[self.regulation_list[2]][0]
        )]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0].format(
                # pylint: disable=no-member
                self.get_origin_program_display(),
                self.origin_program,
                self.get_academic_program_display(),
                self.academic_program,
                self.academic_period_transit
            ) + ', '
        )
        paragraph.add_run(
            self.str_cm[1].format(
                self.council_decision
            )
        )

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
