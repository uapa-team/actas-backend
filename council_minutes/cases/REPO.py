from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, FloatField, EmbeddedDocumentListField
from ..models import Request, Subject
from .case_utils import table_subjects, add_analysis_paragraph

class REPO(Request):
    
    full_name = 'Recurso de reposición'

    reference_id = StringField(requiered=True, max_length=24, min_length=24,
            default='0'*24, display='Id del caso a reponer')
    case_number = StringField(required=True, default='0.0.0', 
            display='Número del caso referido')

    regulation_list = []

    RTA_RATIFICAR = 'RATIFICAR'
    RTA_REPONER = 'REPONER'

    AS_EN_ESPERA = 'EE'
    AS_RATIFICA = 'RT'
    AS_REPONE = 'RP'
    AS_ANULADA = 'AN'
    AS_RENUNCIA = 'RN'
    AS_CHOICES = (
        (AS_EN_ESPERA, 'En espera'),
        (AS_RATIFICA, 'Ratificar'),
        (AS_REPONE, 'Reponer'),
        (AS_ANULADA, 'Anular'),
        (AS_RENUNCIA, 'Desistir'),
    )

    ARCR_EN_ESPERA = 'CEE'
    ARCR_RATIFICA = 'RAT'
    ARCR_REPONE = 'REP'
    ARCR_CHOICES = (
        (ARCR_EN_ESPERA, 'En espera'),
        (ARCR_RATIFICA, 'Ratificar'),
        (ARCR_REPONE, 'Reponer'),
    )

    # @Override
    approval_status = StringField(
        min_length=2, max_length=2, choices=AS_CHOICES,
        default=AS_EN_ESPERA, display='Estado de Aprobación')
    # @Override
    advisor_response = StringField(
        min_length=3, max_length=3, choices=ARCR_CHOICES,
        default=AS_EN_ESPERA, display='Respuesta del Comité')

    str_cm = []

    str_pcm = [
        'en atención al recurso de reposición',
        'decisión del Acta {} de {} en consecuencia, '
    ]

    str_analysis = [
        'Se interpone recurso de reposición sobre la decisión del acta {} de {}, caso {}.',
        'Recomendación del Comité Asesor en acta de comité {} de {} con Concepto: '
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        target = Request.get_case_by_id(self.reference_id)

        paragraph.add_run(self.str_comittee_header + ' ')
        paragraph.add_run(self.str_pcm[0] + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            ' ' + self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_pcm[1].format(target.consecutive_minute, target.year))

        if self.advisor_response == self.ARCR_REPONE:
            if target.advisor_response == Request.ARCR_APROBAR:
                target.advisor_response = Request.ARCR_NO_APROBAR
            elif target.advisor_response == Request.ARCR_NO_APROBAR:
                target.advisor_response = Request.ARCR_APROBAR

        target.pcm_answer(paragraph)

    def pcm_analysis(self, docx):
        analysis_list = []
        target = Request.get_case_by_id(self.reference_id)
        
        analysis_list.append(self.str_analysis[0].format(
            target.consecutive_minute, target.year, self.case_number
        ))

        text = self.str_analysis[1].format(target.consecutive_minute, target.year)
        analysis_list.append(text)
        add_analysis_paragraph(docx, analysis_list + self.extra_analysis)

        para = docx.paragraphs[-1-len(self.extra_analysis)]
        target.pcm_answer(para)