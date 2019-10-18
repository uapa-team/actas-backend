from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, FloatField, EmbeddedDocumentListField, BooleanField, DateField
from ..models import Request, Subject
from .case_utils import table_subjects, add_analysis_paragraph, num_to_month


class TGRA(Request):

    full_name = 'Beca exensión derechos académicos'

    gpa = FloatField(display='Promedio obtenido el semestre anterior')
    gpa_period = StringField(
        display='Periodo en el que se obtiene el promedio')
    target_period = StringField(display='Periodo en el que aplica la beca')

    regulation_list = ['2|2012|CFA']  # List of regulations

    str_case = [
        'Se obtiene el promedio {}, en el periodo académico {} y se solicita la beca para el peri' +
        'odo académico {}.',
        'La coordinación curricular del programa presenta como beneficiario de la BECA EXENCIÓN D' +
        'E DERECHOS ACADÉMICOS del Acuerdo 2 de 2012 de Consejo de Facultad, por obtener el prome' +
        'dio académico ponderado más alto del semestre en las asignaturas cursadas durante el per' +
        'iodo académico inmediatamente anterior.',
        ' la BECA EXENCIÓN DE DERECHOS ACADÉMICOS en el programa de {} ({}) en el periodo {} y ot' +
        'orgar la exención del 100% de derechos académicos por obtener el promedio académico pond' +
        'erado más alto del semestre en las asignaturas cursadas durante el periodo académico inm' +
        'ediatamente anterior.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_case[2].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.target_period
        ))
