from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import EmbeddedDocumentListField, StringField
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph, table_english, get_academic_program


class HOID(Request):

    full_name = 'Homologación de Idioma'

    CT_EXAMEN = 'EX'
    CT_CURSO_NAL = 'CN'
    CT_CURSO_INAL = 'CI'
    CT_CHOICES = (
        (CT_EXAMEN, 'Examen'),
        (CT_CURSO_NAL, 'Curso Nacional'),
        (CT_CURSO_INAL, 'Curso Internacional')
    )

    min_grade = StringField(required=True, default='B1',
                            display='Nivel Requerido')
    certification_type = StringField(
        required=True, choices=CT_CHOICES, display='Tipo de certificación', default=CT_EXAMEN)
    institution = StringField(
        required=True, display='Institución/Examen', default='')
    grade_got = StringField(required=True, default='B1',
                            display='Nivel Obtenido')

    subs = []
    subs.append(Subject(name="Inglés I- Semestral",   code='1000044', credits=3, 
                            tipology=Subject.TIP_PRE_NIVELACION))
    subs.append(Subject(name="Inglés II- Semestral",  code='1000045', credits=3, 
                            tipology=Subject.TIP_PRE_NIVELACION))
    subs.append(Subject(name="Inglés III- Semestral", code='1000046', credits=3, 
                            tipology=Subject.TIP_PRE_NIVELACION))
    subs.append(Subject(name="Inglés IV- Semestral",  code='1000047', credits=3, 
                            tipology=Subject.TIP_PRE_NIVELACION))

    subjects = EmbeddedDocumentListField(
        Subject,
        display='Asignaturas Homologadas',
        default=list(subs)
    )

    regulation_list = ['102|2013|CSU', '001|2016|VAC']

    str_cm = []

    str_pcm = ['Alcanzó el nivel {} en el {}.']

    str_ans = [
        'homologar en el periodo académico {}, el requisito de idioma inglés, ',
        'por obtener una calificación de {} en el exámen {}, siendo {} el mínimo exigido.',
        'teniendo en cuenta que presenta un certificado de estudios expedido por una ' +
        'institución de educación superior {}, indicando que ha cursado un total ' +
        'acumulado de horas equivalente al requerido para alcanzar el nivel {} (375 horas).'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        if self.is_affirmative_response_approval_status():
            self.add_subjects(docx)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        self.add_answer(paragraph)

    def pcm(self, docx):
        final_analysis = []
        analysis = self.str_pcm[0].format(self.grade_got, self.institution)
        final_analysis.append(analysis)
        if len(self.extra_analysis) > 0:
            for i in self.extra_analysis:
                final_analysis.append(i)
        add_analysis_paragraph(docx, final_analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ':\n').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        if self.is_affirmative_response_advisor_response() or self.is_waiting_response_advisor_response():
            self.add_subjects(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        self.add_answer(paragraph)

    def add_answer(self, paragraph):
        paragraph.add_run(self.str_ans[0].format(self.academic_period))
        if self.certification_type == self.CT_EXAMEN:
            paragraph.add_run(self.str_ans[1].format(
                self.grade_got, self.institution, self.min_grade
            ))
        else:
            mod = 'nacional' if self.certification_type == self.CT_CURSO_NAL else 'INTERNACIONAL'
            paragraph.add_run(self.str_ans[2].format(mod, self.grade_got))

    def add_subjects(self, docx):
        data = Subject.subjects_to_array(self.subjects)
        details = []
        details.append(self.certification_type)
        details.append(self.grade_got)
        print(self.student_name)
        details.append(self.student_name)
        details.append(self.student_dni)
        details.append(get_academic_program(self.academic_program))
        details.append(self.academic_program)
        table_english(docx, data, details)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)

