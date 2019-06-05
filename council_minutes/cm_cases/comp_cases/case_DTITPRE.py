from docx import Document
from ...models import Request
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from num2words import num2words  ##pip install num2words
from docx.shared import Pt
from .case_REINPRE import REINPRE
from docx.shared import Cm, Inches


class DTITPRE():

    @staticmethod
    def GET_PLAN(cod_plan):
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == cod_plan:
                large_program = p[1]
                break
        return large_program
    
    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO(request, docx):
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El Consejo de Facultad ')
        if request.approval_status == 'AP':
            DTITPRE.case_DOBLE_TITULACION_PREGRADO_AP(request, docx, para)
        else:
            DTITPRE.case_DOBLE_TITULACION_PREGRADO_NA(request, docx, para)
        
    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_AP(request, docx, paragraph):
        paragraph.add_run('APRUEBA').font.bold = True
        paragraph.add_run(' recomendar al Consejo de Sede que formalice la admisión y ubicación')
        paragraph.add_run('en el programa de pregrado {} – {}, '.format(request.get_academic_program_display(), request.academic_program))
        paragraph.add_run('teniendo en cuenta que el estudiante cuenta con un cupo de créditos suficiente para culminar')
        paragraph.add_run('el segundo plan de estudios. (Acuerdo 155 de 2014 del Consejo Superior Universitario).\n')
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLES(request, docx, paragraph)

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_NA(request, docx, paragraph):
        paragraph.add_run('NO APRUEBA').font.bold = True

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLES(request, docx, paragraph):
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run('1. Datos Generales')
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_DATOS_PERSONALES(request, docx)
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('2. Información Académica:')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_INFORMACION_ACADEMICA(request, docx)
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('3. Cuadro equivalencia y convalidaciones de asignaturas cursadas y aprobadas hasta' 
        + 'la fecha de presentación de la solicitud por parte del estudiante.')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_EQUIVALENCIAS_CONVALIDACIONES(request, docx)
    
    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_DATOS_PERSONALES(request, docx):
        table = docx.add_table(rows=8, cols=3, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        # table.allow_autofit = False
        for cell in table.columns[0].cells:
            cell.width = 400000
        for cell in table.columns[1].cells:
            cell.width = 2400000
        for cell in table.columns[2].cells:
            cell.width = 2400000
        cellp = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('DOBLE TITULACIÓN\n').font.bold = True
        cellp.add_run('Normativa Asociada: Articulo 47 al 50 del Acuerdo 008 de 2008 del CSU y Acuerdo 155 de 2014 del CSU')
        table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(4, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(5, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(6, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(7, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 0).paragraphs[0].add_run('1').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Estudiante')
        table.cell(1, 2).paragraphs[0].add_run(request.student_name)
        table.cell(2, 0).paragraphs[0].add_run('2').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('DNI')
        table.cell(2, 2).paragraphs[0].add_run(request.student_dni)
        table.cell(3, 0).paragraphs[0].add_run('3').font.bold = True
        table.cell(3, 1).paragraphs[0].add_run('Plan de estudios origen (1er plan) – Sede')
        table.cell(3, 2).paragraphs[0].add_run(DTITPRE.GET_PLAN(request.detail_cm['datos_generales']['plan_origen']))
        table.cell(4, 0).paragraphs[0].add_run('4').font.bold = True
        table.cell(4, 1).paragraphs[0].add_run('Código del plan de estudios')
        table.cell(4, 2).paragraphs[0].add_run(request.detail_cm['datos_generales']['plan_origen'])
        table.cell(5, 0).paragraphs[0].add_run('5').font.bold = True
        table.cell(5, 1).paragraphs[0].add_run('Plan de estudios doble titulación (2do plan)')
        table.cell(5, 2).paragraphs[0].add_run(request.get_academic_program_display())
        table.cell(6, 0).paragraphs[0].add_run('6').font.bold = True
        table.cell(6, 1).paragraphs[0].add_run('Código del plan de estudios doble titulación')
        table.cell(6, 2).paragraphs[0].add_run(request.academic_program)
        table.cell(7, 0).paragraphs[0].add_run('7').font.bold = True
        table.cell(7, 1).paragraphs[0].add_run('Fecha de la Solicitud a través del SIA')
        table.cell(7, 2).paragraphs[0].add_run(str(request.detail_cm['datos_generales']['fecha_solicitud'].day) + REINPRE.num_to_month(request.detail_cm['datos_generales']['fecha_solicitud'].month) + str(request.detail_cm['datos_generales']['fecha_solicitud'].year))

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_INFORMACION_ACADEMICA(request, docx):
        table = docx.add_table(rows=4, cols=5, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        for cell in table.columns[0].cells:
            cell.width = 4000000
        for cell in table.columns[1].cells:
            cell.width = 300000
        for cell in table.columns[2].cells:
            cell.width = 300000
        for cell in table.columns[3].cells:
            cell.width = 300000
        for cell in table.columns[4].cells:
            cell.width = 300000
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 0).paragraphs[0].add_run('¿Tuvo calidad de estudiante en el 2do plan?')
        table.cell(0, 1).paragraphs[0].add_run('Si')
        table.cell(0, 3).paragraphs[0].add_run('No')
        if request.detail_cm['informacion_academica']['calidad_estudiante_seg_plan'] == 'Si':
            table.cell(0, 2).paragraphs[0].add_run('X')
        else:
            table.cell(0, 4).paragraphs[0].add_run('X')
        
        table.cell(1, 0).paragraphs[0].add_run('Se encuentra matriculado al momento de presentar la solicitud')
        table.cell(1, 1).paragraphs[0].add_run('Si')
        table.cell(1, 3).paragraphs[0].add_run('No')
        if request.detail_cm['informacion_academica']['matriculado_solicitud'] == 'Si':
            table.cell(1, 2).paragraphs[0].add_run('X')
        else:
            table.cell(1, 4).paragraphs[0].add_run('X')
        
        table.cell(2, 0).paragraphs[0].add_run('PAPA en el primer plan de estudio')
        cellp = table.cell(2, 1).merge(table.cell(2, 4)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run(request.detail_cm['informacion_academica']['PAPA_plan_origen'])
        table.cell(3, 0).paragraphs[0].add_run('Créditos estudio doble titulación')
        cellp = table.cell(3, 1).merge(table.cell(3, 4)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run(request.detail_cm['informacion_academica']['cred_estudio_doble_titulacion'])
    
    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_EQUIVALENCIAS_CONVALIDACIONES(request, docx):
        table = docx.add_table(rows=len(request.detail_cm['equivalencias_convalidaciones']['T_B'])+2, cols=9, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # for column in table.columns:
        #     column.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        # cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('PLAN DE ESTUDIOS (1)').font.bold = True
        cellp = table.cell(0, 2).merge(table.cell(0, 8)).paragraphs[0]
        cellp.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        # cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('PLAN DE ESTUDIOS (2)\nCOMPONENTE DE FUNDAMENTACIÓN (TIPOLOGÍA B)').font.bold = True
        table.cell(1, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 3).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 4).paragraphs[0].add_run('Ob*').font.bold = True
        table.cell(1, 5).paragraphs[0].add_run('Op*').font.bold = True
        table.cell(1, 6).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(1, 7).paragraphs[0].add_run('C*').font.bold = True
        table.cell(1, 8).paragraphs[0].add_run('Nota').font.bold = True
        row = 2
        column = 0
        # table.cell(2, 0).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][0]['codigo_origen'])

        for i in range (0, len(request.detail_cm['equivalencias_convalidaciones']['T_B']) - 1):
            table.cell(row, column).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['codigo_origen'])
            table.cell(row, column+1).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['asignatura_origen'])
            table.cell(row, column+2).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['codigo_seg'])
            table.cell(row, column+3).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['asignatura_seg'])
            if request.detail_cm['equivalencias_convalidaciones']['T_B'][0]['Ob/Op'] == 'Ob':
                table.cell(row, column+4).paragraphs[0].add_run('X')
            else:
                table.cell(row, column+5).paragraphs[0].add_run('X')
            table.cell(row, column+6).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['agrupacion'])
            table.cell(row, column+7).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['C'])
            table.cell(row, column+8).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['nota'])
            row = row + 1
            column = 0


