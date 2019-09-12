from num2words import num2words
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from ...models import Request
from ...pre_cm_cases.comp_cases.case_utils import table_approvals


class simple():

    @staticmethod
    def case_RECURSO_DE_APELACION(request, docx, redirected=False):
        from ..spliter import CasesSpliter
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El Consejo de Facultad en atención al recurso de')
        para.add_run(' apelación, ')
        large_app_st = ''
        for app_st in Request.APPROVAL_STATUS_CHOICES:
            if app_st[0] == request.approval_status:
                large_app_st = app_st[1]
                break
        ap_st = para.add_run(large_app_st)
        ap_st.font.bold = True
        ap_st.font.all_caps = True
        para.add_run(' la decisión del Acta ')
        para.add_run(request.detail_cm['acta_ref'])
        para.add_run(', en consecuencia, ')
        modified_request = request
        modified_request.approval_status = request.detail_cm['approval_status']
        modified_request.type = request.detail_cm['origin_case']
        spliter = CasesSpliter()
        print(modified_request.approval_status)
        spliter.request_case(modified_request, docx, True)

    @staticmethod
    def case_RECURSO_DE_REPOSICION(request, docx, redirected=False):
        from ..spliter import CasesSpliter
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El Consejo de Facultad en atención al recurso de')
        para.add_run(' reposición, ')
        large_app_st = ''
        for app_st in Request.APPROVAL_STATUS_CHOICES:
            if app_st[0] == request.approval_status:
                large_app_st = app_st[1]
                break
        ap_st = para.add_run(large_app_st)
        ap_st.font.bold = True
        ap_st.font.all_caps = True
        para.add_run(' la decisión del Acta ')
        para.add_run(request.detail_cm['acta_ref'])
        para.add_run(', en consecuencia, ')
        modified_request = request
        modified_request.approval_status = request.detail_cm['approval_status']
        modified_request.type = request.detail_cm['origin_case']
        spliter = CasesSpliter()
        print(modified_request.approval_status)
        spliter.request_case(modified_request, docx, True)

    @staticmethod
    def case_RECURSO_DE_REPOSICION_CON_SUBSIDIO_DE_APELACION(request, docx, redirected=False):
        from ..spliter import CasesSpliter
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El Consejo de Facultad en atención al recurso de')
        para.add_run(' reposición con subsidio de apelación, ')
        large_app_st = ''
        for app_st in Request.APPROVAL_STATUS_CHOICES:
            if app_st[0] == request.approval_status:
                large_app_st = app_st[1]
                break
        ap_st = para.add_run(large_app_st)
        ap_st.font.bold = True
        ap_st.font.all_caps = True
        para.add_run(' la decisión del Acta ')
        para.add_run(request.detail_cm['acta_ref'])
        para.add_run(', en consecuencia, ')
        modified_request = request
        modified_request.approval_status = request.detail_cm['approval_status']
        modified_request.type = request.detail_cm['origin_case']
        spliter = CasesSpliter()
        print(modified_request.approval_status)
        spliter.request_case(modified_request, docx, True)

    @staticmethod
    def case_CANCELACION_DE_PERIODO_ACADEMICO_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('cancelar la totalidad de las asignaturas en el periodo ')
        para.add_run(request.detail_cm['period_cancel'])
        para.add_run(', en el programa de ' + large_program)
        para.add_run(' teniendo en cuenta que ')
        para.add_run(request.justification)
        para.add_run(' (Artículo 18 del Acuerdo 008 del Consejo Superior Universitario).')

    @staticmethod
    def case_CANCELACION_DE_PERIODO_ACADEMICO_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run(' cancelar el periodo académico ' +
                     request.academic_period)
        if request.approval_status == 'AP':
            para.add_run(
                ', debido a que justifica documentalmente la fuerza mayor o caso fortuito.')
        else:
            para.add_run(', debido a que ' + request.justification)
        para.add_run(
            ' (Artículo 18 del Acuerdo 008 del Consejo Superior Universitario).')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @staticmethod
    def case_CAMBIO_DE_PERFIL_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
        else:
            para.add_run('NO APRUEBA').font.bold = True
        para.add_run(' el traslado del plan de estudios de ')
        para.add_run(request.detail_cm['from_node'])
        para.add_run(' al plan de estudios de ')
        para.add_run(request.detail_cm['to_node'])
        para.add_run(' de ' + large_program + ' debido a que ')
        para.add_run(request.justification + '.')

    @staticmethod
    def case_AMPLIACION_DE_LA_FECHA_DE_PAGO_DE_DERECHOS_ACADEMICOS_POSGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
        else:
            para.add_run('NO APRUEBA').font.bold = True
        para.add_run(
            ' presentar con concepto positivo al Comité de Matriculas de la Sede')
        para.add_run(
            ' Bogotá, la expedición de un único recibo correspondiente a los')
        para.add_run(
            ' derechos académicos y administrativos para el periodo académico ')
        para.add_run(request.academic_period + ' debido a que ')
        para.add_run(request.justification + '.')

    @staticmethod
    def case_REEMBOLSO_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run('reembolsar {} créditos al estudiante, debido a que {}'.format(
                request.detail_cm['credits'], request.justification))

    @staticmethod
    def case_EXENCION_DE_PAGO_POR_CURSAR_TESIS_COMO_UNICA_ACTIVIDAD_ACADEMICA_POSGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        common = 'pago de {} puntos por derechos académicos en el periodo académico {},'
        common = common + ' condicionado a la inscripción de trabajo final de {} como única '
        common = common + 'actividad académica en el periodo {}'
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            # ¿Los dos periodos mencionados siempre son iguales?
            para.add_run(common.format(
                request.detail_cm['points'], request.academic_period,
                request.get_academic_program_display(), request.academic_period))
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common.format(request.detail_cm['points'],
                                       request.academic_period,
                                       request.get_academic_program_display(),
                                       request.academic_period, request.justification))
            para.add_run(', debido a que {}'.format(request.justification))
        para.add_run('.')

    @staticmethod
    def case_GENERACION_DE_RECIBO_UNICO_DE_PAGO_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        common = 'presentar con concepto {} al Comité de Matrículas de la Sede Bogotá,'
        common = common + ' la expedición de un único recibo correspondiente a saldos pendientes '
        common = common + 'de matrícula para el periodo académico {}'
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            #No se si siempre es "porque se justifica debidamente la solicitud"
            #o si existen más casos
            para.add_run(common.format(request.detail_cm['concept'], request.academic_period))
            para.add_run(', porque justifica debidamente la solicitud. {}'.format(
                request.observation))
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common.format(request.detail_cm['concept'], request.academic_period))
            para.add_run(', debido a que {}.'.format(request.justification))

    @staticmethod
    def case_EXENCION_DE_PAGO_POR_CREDITOS_SOBRANTES_DE_PREGRADO_POSGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
            para.add_run(' otorgar excención del pago de ' +
                         request.detail_cm['points'])
            para.add_run(
                ' puntos de Derechos Académicos, a partir del periodo ' + request.academic_period)
            para.add_run(
                ', y durante el siguiente periodo académico,' +
                ' por tener créditos disponibles al finalizar ')
            para.add_run('estudios del programa de pregrado ' +
                         request.detail_cm['program'] + ', Sede ')
            para.add_run(
                request.detail_cm['campus'] + '. El cálculo de los ' +
                'créditos disponibles se realiza con base')
            para.add_run(
                ' en el cupo de créditos establecido en el Artículo ' +
                '2 del acuerdo 014 de 2008 del Consejo Académico. ')
        else:
            para.add_run('NO APRUEBA').font.bold = True
            para.add_run(
                ' otorgar excención del pago de  Derechos Académicos ' +
                'a partir del periodo ' + request.academic_period)
            para.add_run(
                ', por tener créditos disponibles al finalizar estudios ' +
                'en el programa de pregrado de ')
            para.add_run(request.detail_cm['program'] + ', Sede ' +
                         request.detail_cm['campus'] + ' porque ' + request.justification)
            para.add_run(
                '. (Artículo 58 del acuerdo 008 de 2008 del Consejo Superior Universitario. ')

    @staticmethod
    def case_DEVOLUCION_PROPORCIONAL_DEL_VALOR_PAGADO_POR_CONCEPTO_DE_DERECHOS_MATRICULA_PREGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
            para.add_run(' la devolución proporcional del ' +
                         request.detail_cm['percentage'] + '%')
            para.add_run(
                ' del valor pagado por concepto de derechos de matricula del periodo ')
            para.add_run(request.academic_period +
                         ', teniendo en cuenta la fecha de presentación de la solicitud y que le ')
            para.add_run(
                'fue aprobada la cancelación de periodo en Acta ' + request.detail_cm['acta'])
            para.add_run(
                ' de Consejo de Facultad. (Acuerdo 032 de 2010 del ' +
                'Consejo Superior Universitario, Artículo 1 ')
            para.add_run(' Resolución 1416 de 2013 de Rectoría). ')
        else:
            para.add_run('NO APRUEBA').font.bold = True
            para.add_run(' la devolución proporcional del')
            para.add_run(
                ' valor pagado por concepto de derechos de matricula del periodo ')
            para.add_run(request.academic_period +
                         ', teniendo en cuenta que no le fue aprobada la cancelación de ')
            para.add_run('periodo, según Acta ' + request.detail_cm['acta'])
            para.add_run(' de Consejo de Facultad. (Artículo 1 ')
            para.add_run(' Resolución 1416 de 2013 de Rectoría). ')

    @staticmethod
    def case_DEVOLUCION_DE_CREDITOS_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
        else:
            para.add_run('NO APRUEBA').font.bold = True
        para.add_run(
            ' reintegrar al cupo, los créditos descontados por la ' +
            'cancelación de la(s) sugiente(s) asignatura(s) ')
        para.add_run('en el periodo académico ' + request.academic_period)
        para.add_run(
            '. (Circular 001 de 2019 de Vicerrectoría de Sede Bogotá, Acuerdo 230 de 2016 de ' +
            'Consejo Superior Universitario).')
        table = docx.add_table(
            rows=len(request.detail_cm['subjects'])+2, cols=3, style='Table Grid')
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in table.columns[0].cells:
            cell.width = 915000
        for cell in table.columns[1].cells:
            cell.width = 3620000
        for cell in table.columns[2].cells:
            cell.width = 915000
        table.cell(0, 0).paragraphs[0].add_run('Código SIA').font.bold = True
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 1).paragraphs[0].add_run(
            'Nombre Asignatura').font.bold = True
        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 2).paragraphs[0].add_run('Créditos').font.bold = True
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        index = 1
        credits_sum = 0
        for subject in request.detail_cm['subjects']:
            credits_sum = credits_sum+int(subject['credits'])
            table.cell(index, 0).paragraphs[0].add_run(subject['code'])
            table.cell(index, 1).paragraphs[0].add_run(subject['name'])
            table.cell(index, 2).paragraphs[0].add_run(subject['credits'])
            index = index + 1
        table.cell(index, 2).paragraphs[0].add_run(str(credits_sum))
        cellp = table.cell(index, 0).merge(table.cell(index, 1)).paragraphs[0]
        cellp.add_run('Total Créditos').font.bold = True
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in table.columns[0].cells:
            cell.width = 1500000
        for cell in table.columns[1].cells:
            cell.width = 3050000
        for cell in table.columns[2].cells:
            cell.width = 900000
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def case_ELIMINACION_DE_LA_HISTORIA_ACADEMICA_BAPI_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
            para.add_run(
                ' eliminar la historia académica BAPI, debido a que ' + request.justification + '.')
        else:
            para.add_run('NO APRUEBA').font.bold = True
            para.add_run(
                ' eliminar la historia académica BAPI, debido a que ' + request.justification+'.')

    @staticmethod
    def case_RESERVA_DE_CUPO_ADICIONAL_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run(
            request.detail_cm['index'] + ' reserva de cupo adicional en el periodo académico ')
        para.add_run(request.academic_period)
        if request.approval_status == 'AP':
            para.add_run(', debido a que justifica debidamente la solicitud.')
        else:
            para.add_run(', debido a que ' + request.justification)
        para.add_run(
            ' (Artículo 20 del Acuerdo 008 de 2008 del Consejo Superior Universitario.)')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @staticmethod
    def case_REEMBOLSO_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('devolución proporcional del ' +
                     num2words(float(request.detail_cm['percentage']), lang='es'))
        para.add_run(' por ciento (' + request.detail_cm['percentage'] + '%) ')
        para.add_run(
            'del valor pagado por concepto de derechos de matrícula del periodo ')
        para.add_run(request.detail_cm['period_cancel'])
        para.add_run(
            ', teniendo en cuenta la fecha de presentación de la solicitud y que le fue ' +
            'aprobada la cancelación de periodo en Acta 0')
        para.add_run(request.detail_cm['acta_n'])
        para.add_run(' de ' + request.detail_cm['acta_y'])
        para.add_run(
            ' de Consejo de Facultad. (Acuerdo 032 de 2010 del Consejo Superior ')
        para.add_run(
            'Universitario, Artículo 1 Resolución 1416 de 2013 de Rectoría).')

    @staticmethod
    def case_ADMISION_AUTOMATICA_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('admisión automática al programa ' + large_program)
        para.add_run(' a partir del periodo académico ')
        para.add_run(request.detail_cm['ing_period'])
        para.add_run(
            '. (Acuerdo 070 de 2009 de Consejo Académico y literal c, Artículo 57 del Acuerdo' +
            ' 008 de 2008 del Consejo Superior Universitario.).')

    @staticmethod
    def case_REGISTRO_DE_CALIFICACION_DE_MOVILIDAD_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('calificar ')
        if request.detail_cm['cal'] == 'AP':
            para.add_run('aprobada (AP) ')
        else:
            para.add_run('no aprobada (NA) ')
        para.add_run('la asignatura ' + request.detail_cm['cod_assig'] + ' - ')
        para.add_run(request.detail_cm['nomb_assig'] + ' en el periodo ')
        para.add_run(request.detail_cm['per_assig'] + '.')

    @staticmethod
    def case_DESIGNACION_DE_JURADOS_CALIFICADORES_DE_TESIS_TRABAJO_FINAL_POSGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('designar en el jurado calificador de')
        if request.detail_cm['tesis_trabajo'] == 'Trabajo Final':
            para.add_run('l Trabajo Final de ' +
                         request.detail_cm['nivel_pos'] + ' de ')
        elif request.detail_cm['tesis_trabajo'] == 'Tesis':
            para.add_run(' la Tesis de ' +
                         request.detail_cm['nivel_pos'] + ' de ')
        para.add_run(large_program)
        para.add_run(', cuyo título es: "')
        para.add_run(request.detail_cm['tittle']).font.italic = True
        para.add_run('", a los docentes ')
        count = len(request.detail_cm['doc'])
        for doc in request.detail_cm['doc']:
            count = count - 1
            para.add_run(doc['nomb'])
            if 'dep' in doc:
                para.add_run(
                    ' de la Universidad Nacional de Colombia de la dependencia: ')
                para.add_run(doc['dep'])
            elif 'univ' in doc:
                para.add_run(' de la ' + doc['univ'])
            else:
                raise AttributeError
            if count == 0:
                para.add_run('.')
                return
            elif count == 1:
                para.add_run(' y ')
                break
            else:
                para.add_run(', ')
        length = len(request.detail_cm['doc']) - 1
        para.add_run(request.detail_cm['doc'][length]['nomb'])
        if 'dep' in request.detail_cm['doc'][length]:
            para.add_run(
                ' de la Universidad Nacional de Colombia de la dependencia: ')
            para.add_run(request.detail_cm['doc'][length]['dep'])
        elif 'univ' in request.detail_cm['doc'][length]:
            para.add_run(' de la ' + request.detail_cm['doc'][length]['univ'])
            para.add_run('.')
        else:
            raise AttributeError

    @staticmethod
    def case_MODIFICACION_DE_OBJETIVOS_DE_TESIS_PROPUESTA_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        common = 'cambiar objetivos de Tesis de {} a: “{}”'
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            para.add_run(
                common.format(request.get_academic_program_display(), request.detail_cm['title']))
            if request.observation:
                para.add_run(', ' + request.observation)
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common.format(
                request.get_academic_program_display(), request.detail_cm['title']))
            para.add_run(', debido a que {}'.format(request.justification))
        para.add_run('.')

    @staticmethod
    def case_CARGA_INFERIOR_A_LA_MINIMA_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
        else:
            para.add_run('NO APRUEBA').font.bold = True
        para.add_run(' cursar el periodo académico ' + request.academic_period)
        para.add_run(' con un número de créditos inferior al mínimo exigido, debido a que ')
        if request.approval_status == 'AP':
            para.add_run(' justifica debidamente la solitud.')
        else:
            para.add_run(request.justification)
        para.add_run(' (Artículo 10 del Acuerdo 008 de 2008 del Consejo Superior Universitario).')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @staticmethod
    def case_RETIRO_DEFINITIVO_DEL_PROGRAMA_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.add_run('APRUEBA').font.bold = True
        para.add_run(' presentar con concepto positivo a la ' +
                     'División de Registro y Matrícula, el retiro ')
        para.add_run('voluntario del programa ' + request.get_academic_program_display() +
                     ' (' + request.academic_program + ').')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #No encuentro ningún caso en el que se presente un concepto
        #negativo para este caso o que no se apruebe

    @staticmethod
    def case_CREDITOS_EXCEDENTES_MAPI_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('trasladar ' + request.detail_cm['credits'] +' crédito(s) aprobado(s) en ')
        para.add_run(request.detail_cm['program'])
        if request.approval_status != 'AP':
            para.add_run(' debido a que ' + request.justification)
            return
        para.add_run(' exigidos por la asignatura Trabajo de Grado, el cual se asumirá ' +
                     'como crédito inscrito y aprobado del ')
        para.add_run('componente de libre elección, si en este componente aún hay créditos ' +
                     'por ser aprobados. ')
        para.add_run('(Artículo 16 del Acuerdo 026 de 2012 del Consejo Académico)')

    @staticmethod
    def case_CAMBIO_DE_TIPOLOGIA_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('cambiar de componente la(s) siguiente(s) asignatura(s) del programa ' +
                     large_program)
        if request.approval_status == 'AP':
            para.add_run(', cursada en el periodo académico ')
            para.add_run(request.detail_cm['periodo'] + ' así:')
        else:
            para.add_run(', debido a que ' +request.justification +'.')
        table = docx.add_table(rows=len(request.detail_cm['subjects'])+1, cols=5)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)

        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[0].width = 700000
        # table.columns[1].width = 2000000
        # table.columns[2].width = 600000
        # table.columns[3].width = 1050000
        # table.columns[4].width = 1050000
        for cell in table.columns[0].cells:
            cell.width = 750000
        for cell in table.columns[1].cells:
            cell.width = 2000000
        for cell in table.columns[2].cells:
            cell.width = 600000
        for cell in table.columns[3].cells:
            cell.width = 1050000
        for cell in table.columns[4].cells:
            cell.width = 1050000
        table.cell(0, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(0, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(0, 2).paragraphs[0].add_run('Nota').font.bold = True
        table.cell(0, 3).paragraphs[0].add_run('Componente Registrado').font.bold = True
        if request.approval_status == 'AP':
            table.cell(0, 4).paragraphs[0].add_run('Nuevo Componente').font.bold = True
        else:
            table.cell(0, 4).paragraphs[0].add_run('Componente Solicitado').font.bold = True
        index = 0
        for subject in request.detail_cm['subjects']:
            table.cell(index+1, 0).paragraphs[0].add_run(subject['cod'])
            table.cell(index+1, 1).paragraphs[0].add_run(subject['subject'])
            table.cell(index+1, 2).paragraphs[0].add_run(subject['nota'])
            table.cell(index+1, 3).paragraphs[0].add_run(subject['to'])
            table.cell(index+1, 4).paragraphs[0].add_run(subject['td'])
            index = index + 1
        para = docx.add_paragraph()

    @staticmethod
    def case_TRANSITO_ENTRE_PROGRAMAS_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        large_program2 = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.detail_cm['origen']:
                large_program2 = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('tránsito del programa ' + large_program2 + ' al programa de ' + large_program)
        if request.approval_status == 'AP':
            para.add_run(', a partir del periodo académico ' + request.detail_cm['desde'] +
                         ' (Artículo 3, Resolución 241 de ')
            para.add_run('2009 de la Vicerrectoría Académica).')
        else:
            para.add_run(' porque ' + request.justification + '.')

    @staticmethod
    def case_CAMBIO_DE_DIRECTIOR_CODIRECTOR_JURADO_O_EVALUADOR_POSGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('cambiar el(la) ' + request.detail_cm['rol'] + '(a) de')
        if request.detail_cm['testra'] == 'Tesis':
            para.add_run(' la Tesis de ')
        else:
            para.add_run('l Trabajo Final de ')
        para.add_run(large_program + ', en reemplazo del(de la) profesor(a) ' +
                     request.detail_cm['antiguo'])

        if request.approval_status == 'AP':
            para.add_run('; designa nuevo(a) ' + request.detail_cm['rol'] +'(a) de')
            if request.detail_cm['testra'] == 'Tesis':
                para.add_run(' la Tesis de ')
            else:
                para.add_run('l Trabajo Final de ')
            para.add_run(large_program + ' cuyo título es: ')
            para.add_run('"' + request.detail_cm['titulo'] + '"').font.italic = True
            para.add_run(' al(a la) profesor(a) ' + request.detail_cm['nuevo'] + ' del '+
                         request.detail_cm['depto'] + '.')
        else:
            para.add_run(' porque ' + request.justification+ '.')

    @staticmethod
    def case_DESIGNACION_DE_CODIRECTOR_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        common = 'designar codirector de Tesis de {} con título “{}” '
        common = common + 'aprobado en el Acta No. {}, al profesor/a {}'
        common = common.format(request.get_academic_program_display(),
                               request.detail_cm['title'],
                               request.detail_cm['minutes_approved'],
                               request.detail_cm['professor_name'])
        if request.detail_cm['professor_faculty']:
            information = ' del Departamento {} de la Facultad de {}'
            information = information.format(request.detail_cm['professor_department'],
                                             request.detail_cm['professor_faculty'])
        if request.detail_cm['professor_university']:
            information = ' de la {}'.format(request.detail_cm['professor_university'])
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            para.add_run(common + information)
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common + information)
            para.add_run(', debido a que {}'.format(request.justification))

    @staticmethod
    def case_EVALUADOR_ADICIONAL_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        common = 'designar evaluador adicional del Trabajo Final de {}, '
        common = common + 'cuyo título es: “{}”, al profesor {}'
        common = common.format(request.get_academic_program_display(),
                               request.detail_cm['title'],
                               request.detail_cm['professor_name'])
        if request.detail_cm['professor_faculty']:
            information = ' del Departamento {} de la Facultad de {}'
            information = information.format(request.detail_cm['professor_department'],
                                             request.detail_cm['professor_faculty'])
        if request.detail_cm['professor_university']:
            information = ' de la {}'.format(request.detail_cm['professor_university'])
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            para.add_run(common + information)
            para.add_run(', quien deberá dirimir la diferencia calificando el trabajo final ' +
                         'como aprobado o reprobado. (Acuerdo 56 de 2012 Consejo Superior ' +
                         'Universitario)')
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common + information)
            para.add_run(', debido a que {}'.format(request.justification))
        para.add_run('.')

    @staticmethod
    def case_TRABAJO_DE_GRADO_PREGADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('inscribir la(s) siguiente(s) asignatura(s) en el periodo académico ')
        para.add_run(request.detail_cm['per_insc'])
        cod_assig = ''
        nom_assig = ''
        if request.detail_cm['modalidad'] == 'pasantia':
            para.add_run(', modalidad pasantía en "')
            para.add_run(request.detail_cm['empresa'])
            cod_assig = '2015289'
            nom_assig = 'Trabajo de grado'
        elif request.detail_cm['modalidad'] == 'trabajo':
            para.add_run(', modalidad trabajo investigativo titulado "')
            para.add_run(request.detail_cm['titulo']).font.italic = True
            cod_assig = '2025990'
            nom_assig = 'Trabajo de grado - Modalidad Trabajos Investigativos'
        else:
            raise AttributeError
        if 'docente' in request.detail_cm:
            para.add_run('", dirigida por el profesor ')
            para.add_run(request.detail_cm['docente'])
        para.add_run(', debido a que ')
        para.add_run(request.justification + '.')
        table = docx.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[0].width = 600000
        # table.columns[1].width = 2800000
        # table.columns[2].width = 600000
        # table.columns[3].width = 600000
        # table.columns[4].width = 600000
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in table.columns[0].cells:
            cell.width = 750000
        for cell in table.columns[1].cells:
            cell.width = 2300000
        for cell in table.columns[2].cells:
            cell.width = 800000
        for cell in table.columns[3].cells:
            cell.width = 800000
        for cell in table.columns[4].cells:
            cell.width = 800000

        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0, 0).paragraphs[0].add_run('Código SIA').font.bold = True
        table.cell(0, 1).paragraphs[0].add_run('Nombre Asignatura').font.bold = True
        table.cell(0, 2).paragraphs[0].add_run('Grupo').font.bold = True
        table.cell(0, 3).paragraphs[0].add_run('T').font.bold = True
        table.cell(0, 4).paragraphs[0].add_run('C').font.bold = True
        table.cell(1, 0).paragraphs[0].add_run(cod_assig)
        table.cell(1, 1).paragraphs[0].add_run(nom_assig)
        table.cell(1, 2).paragraphs[0].add_run(request.detail_cm['group'])
        table.cell(1, 3).paragraphs[0].add_run('P')
        table.cell(1, 4).paragraphs[0].add_run('6')

    @staticmethod
    def case_APROBACION_PROYECTO_PROPUESTA_Y_DESIGNACION_DE_DIRECTOR_POSGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA: ').font.bold = True
        else:
            para.add_run('NO APRUEBA: ').font.bold = True
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('1. Calificación Aprobada (AP) a ' + request.detail_cm['testra'] +
                     ' de ' + large_program + ', cuyo título es: ')
        para.add_run('"' + request.detail_cm['titulo'] + '".').font.italic = True
        if request.approval_status == 'AP':
            para = docx.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('2. Designar director(a) de')
            if request.detail_cm['testra'] == 'Tesis de Maestría':
                para.add_run(' la Tesis de ')
            else:
                para.add_run('l Trabajo Final de ').font.bold = True
            para.add_run(large_program + ' cuyo título es: ')
            para.add_run('"' + request.detail_cm['titulo'] + '"').font.italic = True
            para.add_run(' al(a la) profesor(a) ' + request.detail_cm['prof'] +
                         ' del ' + request.detail_cm['depto'] + '.')

    @staticmethod
    def case_MODIFICACION_DE_JURADOS_CALIFICADORES_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('designar en el jurado calificador de ' + \
             request.detail_cm['subject'] + ' en ')
        para.add_run(request.get_academic_program_display() + ', cuyo título es: "' +
                     request.detail_cm['project_title']+'"')
        para.add_run(', al(los) profesor(es) ')
        for professor in request.detail_cm['professors']:
            para.add_run(professor['name'] + " - " + professor['institution'] +
                         " - " + professor['country'] + ". ")

    @staticmethod
    def case_BECA_MEJOR_PROMEDIO_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA BECA EXCENCIÓN DE DERECHOS ACADÉMICOS ').font.bold = True
        else:
            para.add_run('NO APRUEBA BECA EXCENCIÓN DE DERECHOS ACADÉMICOS ').font.bold = True
        para.add_run('del programa ' + large_program +
                     ' por obtener el promedio académico ponderado más alto del semestre en las ')
        para.add_run('asignaturas cursadas durante el periodo académico inmediatamente anterior')
        if request.approval_status == 'AP':
            para.add_run(', y otorga exención de derechos académicos. ')
        else:
            para.add_run('. ')
        para.add_run('(Artículo 8 Acuerdo 02 de 2012 de Consejo de Facultad). ')

    @staticmethod
    def case_EXCENCION_POR_MEJORES_SABER_PRO_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA BECA EXCENCIÓN DE DERECHOS ACADÉMICOS ').font.bold = True
        else:
            para.add_run('NO APRUEBA BECA EXCENCIÓN DE DERECHOS ACADÉMICOS ').font.bold = True
        para.add_run('del programa ' + large_program + ' por obtener un excelente resultado ' +
                     'en el exámen de estado SABER-PRO')
        if request.approval_status == 'AP':
            para.add_run(' y otorga exención de derechos académicos. ')
        else:
            para.add_run('. ')

    @staticmethod
    def case_REINGRESO_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        common = 'reingreso por única vez en el programa de {}, a partir del periodo {}'
        common = common.format(request.get_academic_program_display(), request.academic_period)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            para.add_run(common + '. El reingreso del estudiante estará regido por'+\
                ' el Acuerdo 008 de 2008 del Consejo Superior Universitario')
            if request.observation:
                para.add_run('. {}'.format(request.observation))
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common + ', debido a que {}'.format(request.justification))
        para.add_run('.')

    @staticmethod
    def case_REGISTRO_DE_CALIFICACION_DEL_PROYECTO_Y_EXAMEN_DOCTORAL_POSGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
        else:
            para.add_run('NO APRUEBA:').font.bold = True
        calificacion = {'AP': 'aprobado', 'NA': 'no aprobado'}
        for i in range(0, len(request.detail_cm['casos'])):
            if request.detail_cm['casos'][i]['caso'] == 'examen':
                Item1 = 'Calificar {} ({})'.format(calificacion[
                    request.detail_cm['casos'][i]['calificacion']],
                                                   request.detail_cm['casos'][i]['calificacion'])
                Item2 = ' el Examen de calificación con código {}'
                Item2 = Item2.format(request.detail_cm['casos'][i]['codigo'])
                Item3 = ' en el periodo académico {}.'.format(request.academic_period)
                Item = Item1 + Item2 + Item3
                para = docx.add_paragraph(Item, style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(0)
            else:
                Item1 = 'Calificar {} ({})'.format(calificacion[
                    request.detail_cm['casos'][i]['calificacion']],
                                                   request.detail_cm['casos'][i]['calificacion'])
                Item2 = ' el Proyecto de Tesis de {} con código {}'.format(
                    request.get_academic_program_display(), request.detail_cm['casos'][i]['codigo'])
                Item3 = ' en el periodo académico {}, cuyo título es: “{}”.'.format(
                    request.academic_period, request.detail_cm['casos'][i]['titulo'])
                Item = Item1 + Item2 + Item3
                para = docx.add_paragraph(Item, style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(0)

    @staticmethod
    def case_CAMBIO_DE_PROYECTO_DE_TESIS(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        common = 'cambiar título de Tesis de {} a: “{}”'.format(
            request.get_academic_program_display(), request.detail_cm['titulo'])
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            para.add_run(common)
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common + ', debido a que {}'.format(request.justification))
        para.add_run('.')

    @staticmethod
    def case_EXPEDICION_DE_RECIBO_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        if request.detail_cm['caso'] == 'nuevo':
            common = 'expedir un nuevo recibo de pago de derechos de matrícula con ' + \
                'cambio de fecha, para el periodo académico {}'.format(request.academic_period)
        if request.detail_cm['caso'] == 'unico':
            common = 'la expedición de un único recibo correspondiente a ' + \
                'los derechos académicos y administrativos para el periodo académico {}'.format(
                    request.academic_period)
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            para.add_run(common)
            if request.observation:
                para.add_run(', {}'.format(request.observation))
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run(common + ', debido a que {}'.format(request.justification))
        para.add_run('.')

    @staticmethod
    def case_INFORME_DE_AVANCE_DE_TESIS_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        common = '{} de Doctorado ({}) en {}, en el periodo {}'.format(
            request.detail_cm['caso'], request.detail_cm['codigo'],
            request.get_academic_program_display(), request.academic_period)
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
            para.add_run('calificación avance satisfactorio (AS) a ' + common +
                         ', teniendo en cuenta que el estudiante entregó el documento' + \
                             ' para nombramiento de jurados')
        else:
            para.add_run('NO APRUEBA ').font.bold = True
            para.add_run('calificación avance satisfactorio (AS) a ' + common + \
                 ', debido a que {}'.format(request.justification))
        para.add_run('.')

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_CON_CARGA_INFERIOR_A_LA_MINIMA_PREGRADO(
            request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
        else:
            para.add_run('NO APRUEBA: ').font.bold = True
        para.paragraph_format.space_after = Pt(0)
        para = docx.add_paragraph(style='List Number')
        para.add_run('Cursar el periodo académico ' + request.academic_period + \
                     ' con un número de ')
        para.add_run('créditos inferior al mínimo exigido, porque ' + request.justification)
        para.add_run('. (Artículo 10 del Acuerdo 008 de 2008 del Consejo Superior Universitario)')
        para.paragraph_format.space_after = Pt(0)
        para = docx.add_paragraph(style='List Number')
        para.add_run('Cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del periodo ')
        para.add_run(request.academic_period + ', porque ' + request.justification)
        para.add_run('. (Artículo 15 del Acuerdo 008 de 2008 del Consejo Superior Universitario)')
        para.paragraph_format.space_after = Pt(0)
        table = docx.add_table(
            rows=len(request.detail_cm['subjects'])+1, cols=5)
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 750000
        table.columns[1].width = 2300000
        table.columns[2].width = 800000
        table.columns[3].width = 800000
        table.columns[4].width = 800000
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in table.columns[0].cells:
            cell.width = 750000
        for cell in table.columns[1].cells:
            cell.width = 2300000
        for cell in table.columns[2].cells:
            cell.width = 800000
        for cell in table.columns[3].cells:
            cell.width = 800000
        for cell in table.columns[4].cells:
            cell.width = 800000
        table.cell(0, 0).paragraphs[0].add_run('Código SIA').font.bold = True
        table.cell(0, 1).paragraphs[0].add_run('Nombre Asignatura').font.bold = True
        table.cell(0, 2).paragraphs[0].add_run('Grupo').font.bold = True
        table.cell(0, 3).paragraphs[0].add_run('Tipología').font.bold = True
        table.cell(0, 4).paragraphs[0].add_run('Créditos').font.bold = True
        index = 0
        for subject in request.detail_cm['subjects']:
            table.cell(index+1, 0).paragraphs[0].add_run(subject['cod_sia'])
            table.cell(index+1, 1).paragraphs[0].add_run(subject['name_asig'])
            table.cell(index+1, 2).paragraphs[0].add_run(subject['group_asig'])
            table.cell(index+1, 3).paragraphs[0].add_run(subject['tipo_asig'])
            table.cell(index+1, 4).paragraphs[0].add_run(subject['creds_asig'])
            index = index + 1

    @staticmethod
    def case_HOMOLOGACION_DE_ASIGNATURAS_INTERCAMBIO_ACADEMICO_INTERNACIONAL_PREGRADO(
            request, docx, redirected=False):
        assign = ['2011183 - Intercambio Académico Internacional',
                  '2014269 - Intercambio Académico Internacional Prórroga',
                  '2026630 - Intercambio Académico Internacional II',
                  '2026631 - Intercambio Académico Internacional II Prórroga']
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
        else:
            para.add_run('NO APRUEBA: ').font.bold = True
        item = 'Homologar las siguientes asignaturas cursadas en Intercambio ' + \
            'Académico Internacional en '+\
            request.detail_cm['inst'] + ' de la siguiente manera ' + \
                '(Artículo 35, Acuerdo 008 de 2008 ' +\
                'del Consejo Superior Universitario y Resolución 105 ' + \
                    'de 2017 de Vicerrectoría Académica):'
        para = docx.add_paragraph(item, style='List Number')
        para.paragraph_format.space_after = Pt(0)
        asign_number = len(request.detail_cm['subjects'])
        tipology = {}
        periods = []
        for asign in request.detail_cm['subjects']:
            if asign['tipo_asig'] in tipology:
                tipology.update({asign['tipo_asig']: tipology[
                    asign['tipo_asig']] + int(asign['creds_asig'])})
            else:
                tipology.update({asign['tipo_asig']: int(asign['creds_asig'])})
            if asign['per_asig'] not in periods:
                periods.append(asign['per_asig'])
        tipology_number = len(tipology.keys())
        table = docx.add_table(rows=(4+asign_number+tipology_number), cols=8, style='Table Grid')
        table.style.font.size = Pt(8)
        table.columns[0].width = 500000
        table.columns[1].width = 550000
        table.columns[2].width = 1600000
        table.columns[3].width = 300000
        table.columns[4].width = 300000
        table.columns[5].width = 400000
        table.columns[6].width = 1400000
        table.columns[7].width = 400000
        for cell in table.columns[0].cells:
            cell.width = 500000
        for cell in table.columns[1].cells:
            cell.width = 550000
        for cell in table.columns[2].cells:
            cell.width = 1600000
        for cell in table.columns[3].cells:
            cell.width = 300000
        for cell in table.columns[4].cells:
            cell.width = 300000
        for cell in table.columns[5].cells:
            cell.width = 400000
        for cell in table.columns[6].cells:
            cell.width = 1400000
        for cell in table.columns[7].cells:
            cell.width = 400000
        cellp = table.cell(0, 0).merge(table.cell(0, 7)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('{}\t\t\tDNI.{}'.format(
            request.student_name, request.student_dni)).font.bold = True
        cellp = table.cell(1, 0).merge(table.cell(1, 5)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('Asignaturas a homologar en el plan de estudios de {} ({})'.format(
            request.get_academic_program_display(), request.academic_program)).font.bold = True
        cellp = table.cell(1, 6).merge(table.cell(1, 7)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('Asignaturas cursadas en {}'.format(
            request.detail_cm['inst'])).font.bold = True
        for i in range(2, asign_number + 3):
            for j in range(8):
                table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 0).paragraphs[0].add_run('Periodo').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('C').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('T').font.bold = True
        table.cell(2, 5).paragraphs[0].add_run('Nota').font.bold = True
        table.cell(2, 6).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 7).paragraphs[0].add_run('Nota').font.bold = True
        count = 3
        for subject in request.detail_cm['subjects']:
            table.cell(count, 0).paragraphs[0].add_run(subject['per_asig'])
            table.cell(count, 1).paragraphs[0].add_run(subject['cod_asig'])
            table.cell(count, 2).paragraphs[0].add_run(subject['new_name_asig'])
            table.cell(count, 3).paragraphs[0].add_run(subject['creds_asig'])
            table.cell(count, 4).paragraphs[0].add_run(subject['tipo_asig'])
            table.cell(count, 5).paragraphs[0].add_run(subject['new_cal_asig'])
            table.cell(count, 6).paragraphs[0].add_run(subject['old_name_asig'])
            table.cell(count, 7).paragraphs[0].add_run(subject['old_cal_asig'])
            count += 1
        total_homologated = 0
        for tip in tipology:
            text = 'Céditos homologados ' + str(tip)
            table.cell(count, 0).merge(table.cell(count, 5)).paragraphs[0].add_run(text)
            table.cell(count, 0).merge(table.cell(count, 5)).paragraphs[
                0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(count, 6).merge(table.cell(count,
                                                  7)).paragraphs[0].add_run(str(tipology[tip]))
            table.cell(count, 6).merge(table.cell(count, 7)).paragraphs[
                0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            total_homologated += int(tipology[tip])
            count += 1
        table.cell(count, 0).merge(table.cell(count, 5)).paragraphs[
            0].add_run('Total créditos que se homologan')
        table.cell(count, 0).merge(table.cell(count, 5)).paragraphs[
            0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(count, 6).merge(table.cell(count,
                                              7)).paragraphs[0].add_run(str(total_homologated))
        table.cell(count, 6).merge(table.cell(count, 7)).paragraphs[
            0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para = docx.add_paragraph(style='List Number')
        final_text = 'Calificar'
        other_count = 0
        for _ in periods:
            final_text += ' aprobada (AP)' \
                if request.detail_cm['califications'][other_count] == 'AP'\
                     else 'no aprobada (NA)'
            final_text += ' la asignatura ' + assign[other_count] + \
                ' en el periodo ' + periods[other_count]
            other_count += 1
            if other_count + 1 == len(periods):
                final_text += ' y'
        para.add_run(final_text)

    @staticmethod
    def case_HOMOLOGACION_DE_ASIGNATURAS_CONVENIO_CON_UNIVERSIDAD_ANDES_PREGRADO(
            request, docx, redirected=False):
        assign = ['2011302 - Asignatura Por Convenio Con Universidad De Los Andes I - Pregrado',
                  '2012698 - Asignatura Por Convenio Con Universidad De Los Andes II - Pregrado']
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
        else:
            para.add_run('NO APRUEBA: ').font.bold = True
        total_creds = 0
        acum_papa = 0
        for subject in request.detail_cm['subjects']:
            total_creds += int(subject['creds_asig'])
            acum_papa += int(subject['creds_asig']) * \
                float(subject['cal_asign'])
        para = docx.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        mini_papa = acum_papa / total_creds
        para.add_run('Calificación ')
        if mini_papa > 3:
            para.add_run('aprobada (AP)')
        else:
            para.add_run('no aprobada (NA)')
        para.add_run(' en la asignatura ')
        para.add_run(assign[int(request.detail_cm['index']) - 1])
        para.add_run(', en el periodo ')
        para.add_run(request.academic_period)
        para.add_run('.')
        para = docx.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Homologar, en el periodo académico ')
        para.add_run(request.academic_period)
        para.add_run(
            ', la(s) siguiente(s) asignatura(s) cursada(s) en el Convenio en la ')
        para.add_run(
            'Universidad de los Andes de la siguiente manera (Artículo 35 de Acuerdo')
        para.add_run(' 008 de 2008 del Consejo Superior Universitario):')
        subjects = []
        details = [request.student_name, request.student_dni,
                   request.academic_program, 'Universidad de los Andes']
        for subject in request.detail_cm['subjects']:
            subjects.append(
                [request.academic_period, subject['cod_asig'], subject['name_asig'],
                 subject['creds_asig'], 'L', subject['cal_asign'], subject['name_asig'],
                 subject['cal_asign']])
        table_approvals(docx, subjects, details)

    @staticmethod
    def case_HOMOLOGACION_DE_ASIGNATURAS_CONVENIO_CON_UNIVERSIDAD_ANDES_POSGRADO(
            request, docx, redirected=False):
        assign = ['2024944 - Asignatura Por Convenio Con Universidad De Los Andes I - Posgrado',
                  '2024945 - Asignatura Por Convenio Con Universidad De Los Andes II - Posgrado']
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
        else:
            para.add_run('NO APRUEBA: ').font.bold = True
        total_creds = 0
        acum_papa = 0
        for subject in request.detail_cm['subjects']:
            total_creds += int(subject['creds_asig'])
            acum_papa += int(subject['creds_asig']) * \
                float(subject['cal_asign'])
        para = docx.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        mini_papa = acum_papa / total_creds
        para.add_run('Calificación ')
        if mini_papa > 3:
            para.add_run('aprobada (AP)')
        else:
            para.add_run('no aprobada (NA)')
        para.add_run(' en la asignatura ')
        para.add_run(assign[int(request.detail_cm['index']) - 1])
        para.add_run(', en el periodo ')
        para.add_run(request.academic_period)
        para.add_run('.')
        para = docx.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Homologar, en el periodo académico ')
        para.add_run(request.academic_period)
        para.add_run(
            ', la(s) siguiente(s) asignatura(s) cursada(s) en el Convenio en la ')
        para.add_run(
            'Universidad de los Andes de la siguiente manera (Artículo 35 de Acuerdo')
        para.add_run(' 008 de 2008 del Consejo Superior Universitario):')
        subjects = []
        details = [request.student_name, request.student_dni,
                   request.academic_program, 'Universidad de los Andes']
        for subject in request.detail_cm['subjects']:
            subjects.append(
                [request.academic_period, subject['cod_asig'], subject['name_asig'],
                 subject['creds_asig'], 'L', subject['cal_asign'], subject['name_asig'],
                 subject['cal_asign']])
        table_approvals(docx, subjects, details)

    @staticmethod
    def case_ACLARACION_DE_DECISION_PREGRADO(request, docx, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_ACLARACION_DE_DECISION_POSGRADO(request, docx, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_ADICION_DE_CODIRECTOR_POSGRADO(request, docx, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_BECA_EXENSION_DERECHOS_ACADEMICOS(request, docx, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_PROYECTO_DE_TESIS_O_TRABAJO_FINAL_DE_MAESTRIA_POSGRADO(
            request, docx, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_PROYECTO_DE_TESIS_DE_DOCTORADO_POSGRADO(request, docx, redirected=False):
        raise NotImplementedError
