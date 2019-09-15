from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt


class HOIDPXX():

    @staticmethod
    def case_HOMOLOGACION_INGLES(request, docx, redirected=False):
        HOIDPXX.case_HOMOLOGACION_INGLES_Analysis(request, docx)
        HOIDPXX.case_HOMOLOGACION_INGLES_Answers(request, docx)

    @staticmethod
    def case_HOMOLOGACION_INGLES_Analysis(request, docx):
        para = docx.add_paragraph()
        para.add_run('Analisis:')
        analysis_para = docx.add_paragraph()
        analysis_para.paragraph_format.left_indent = Pt(36)
        analysis_str = '1. Obtuvo un puntaje de {} en el examen {}'
        analysis = analysis_str.format(
            request['detail_cm']['grade_got'], request['detail_cm']['institution'])
        analysis_para.add_run(analysis)

    @staticmethod
    def case_HOMOLOGACION_INGLES_Answers(request, docx):
        para = docx.add_paragraph()
        run = para.add_run('Concepto: ')
        run.font.bold = True
        answer_para = docx.add_paragraph()
        answer_para.paragraph_format.left_indent = Pt(36)
        answer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        answer_1 = 'El Comité Asesor recomienda al Consejo de Facultad{} homologar,' + \
            ' en el periodo académico {}, el requisito de idioma inglés por obtener' + \
            ' un puntaje/calificación de {} en el examen {}, siendo {} el mínimo exigido. '
        answer_para.add_run(answer_1.format('', request['academic_period'], 
            request['detail_cm']['grade_got'], request['detail_cm']['institution'],
            request['detail_cm']['min_grade']))